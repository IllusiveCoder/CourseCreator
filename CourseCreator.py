import sys
import pdfkit as pdf
from docx import Document
from docx.shared import Pt
from PyQt5.QtWidgets import QApplication, QMainWindow, QTabWidget, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QTextEdit, QPushButton, QSizePolicy, QStyleFactory, QCheckBox, QButtonGroup
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Create Programming Courses")
        self.setStyleSheet("QMainWindow {background-color: #f2f2f2;}")  # Set background color
        self.setFixedSize(400, 300)  # Set fixed window size

        # Haupt-Widget erstellen
        self.mainWidget = QWidget(self)
        self.setCentralWidget(self.mainWidget)

        # Vertikales Layout für das Haupt-Widget erstellen
        layout = QVBoxLayout(self.mainWidget)
        layout.setContentsMargins(20, 20, 20, 20)  # Set margins
        layout.setSpacing(20)  # Set spacing between widgets

        # Tab-Widget erstellen
        self.tabWidget = QTabWidget()
        layout.addWidget(self.tabWidget)

        # Tabs erstellen
        self.createBasicDataTab()
        self.createContentTab()
        self.createFormatTab()

        # Strings zur Speicherung der Eingaben erstellen
        self.authorInputText = ""
        self.titleInputText = ""
        self.topicsInputText = ""
        self.contentInputText = ""

    def createBasicDataTab(self):
        # Tab für "Basic Data" erstellen
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(10)  # Set spacing between widgets

        # Labels und Eingabefelder hinzufügen
        authorLabel = QLabel("Author:")
        authorInput = QTextEdit()
        authorInput.setFixedHeight(30)
        layout.addWidget(authorLabel)
        layout.addWidget(authorInput)

        titleLabel = QLabel("Title:")
        titleInput = QTextEdit()
        titleInput.setFixedHeight(30)
        layout.addWidget(titleLabel)
        layout.addWidget(titleInput)

        topicsLabel = QLabel("Topics:")
        topicsInput = QTextEdit()
        layout.addWidget(topicsLabel)
        layout.addWidget(topicsInput)

        # Buttons hinzufügen
        buttonLayout = QHBoxLayout()
        layout.addLayout(buttonLayout)
        buttonLayout.addStretch()  # Add stretchable space

        quitButton = QPushButton("Quit")
        nextButton = QPushButton("Next")
        buttonLayout.addWidget(quitButton)
        buttonLayout.addWidget(nextButton)

        # Event-Handler für den "Next"-Button hinzufügen
        nextButton.clicked.connect(self.nextTab)

        # Event-Handler für den "Quit"-Button hinzufügen
        quitButton.clicked.connect(self.quitApplication)

        # Signal-Slot-Verbindungen für die Texteingabefelder
        authorInput.textChanged.connect(lambda: self.setAuthorInput(authorInput.toPlainText()))
        titleInput.textChanged.connect(lambda: self.setTitleInput(titleInput.toPlainText()))
        topicsInput.textChanged.connect(lambda: self.setTopicsInput(topicsInput.toPlainText()))

        self.tabWidget.addTab(tab, "Basic Data")

    def createContentTab(self):
        # Tab für "Content" erstellen
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(10)  # Set spacing between widgets

        contentLabel = QLabel("Content:")
        contentInput = QTextEdit()
        layout.addWidget(contentLabel)
        layout.addWidget(contentInput)

        # Buttonshinzufügen
        buttonLayout = QHBoxLayout()
        layout.addLayout(buttonLayout)
        buttonLayout.addStretch()  # Add stretchable space

        backButton = QPushButton("Back")
        nextButton = QPushButton("Next")
        buttonLayout.addWidget(backButton)
        buttonLayout.addWidget(nextButton)

        # Event-Handler für die Buttons hinzufügen
        backButton.clicked.connect(self.backToBasicDataTab)
        nextButton.clicked.connect(self.nextTab)

        # Signal-Slot-Verbindung für das Texteingabefeld
        contentInput.textChanged.connect(lambda: self.setContentInput(contentInput.toPlainText()))

        self.tabWidget.addTab(tab, "Content")

    def createFormatTab(self):
        # Tab für "Formats" erstellen
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(10)  # Set spacing between widgets

        formatLabel = QLabel("Format:")
        formatLabel.setAlignment(Qt.AlignLeft)
        layout.addWidget(formatLabel)

        # Checkboxen hinzufügen
        checkboxLayout = QVBoxLayout()
        layout.addLayout(checkboxLayout)

        # Checkboxen erstellen
        formats = ["PDF", "Markdown", "HTML", "Word Documents"]
        buttonGroup = QButtonGroup(self)
        for format in formats:
            checkbox = QCheckBox(format)
            checkbox.setFixedHeight(30)
            checkboxLayout.addWidget(checkbox)
            buttonGroup.addButton(checkbox)

        # Nur eine Checkbox kann gleichzeitig ausgewählt sein
        buttonGroup.setExclusive(True)

        # Buttons hinzufügen
        buttonLayout = QHBoxLayout()
        layout.addLayout(buttonLayout)
        buttonLayout.addStretch()  # Add stretchable space

        backButton = QPushButton("Back")
        finishButton = QPushButton("Finish")
        buttonLayout.addWidget(backButton)
        buttonLayout.addWidget(finishButton)

        # Event-Handler für die Buttons hinzufügen
        backButton.clicked.connect(self.backToContentTab)
        finishButton.clicked.connect(self.finish)

        self.tabWidget.addTab(tab, "Formats")

    def nextTab(self):
        # Zum nächsten Tab wechseln
        currentIndex = self.tabWidget.currentIndex()
        nextIndex = currentIndex + 1 if currentIndex < self.tabWidget.count() - 1 else 0
        self.tabWidget.setCurrentIndex(nextIndex)

    def backToBasicDataTab(self):
        # Zum ersten Tab "Basic Data" zurückkehren
        self.tabWidget.setCurrentIndex(0)

    def backToContentTab(self):
        # Zum zweiten Tab "Content" zurückkehren
        self.tabWidget.setCurrentIndex(1)

    def createHtml(self):
    # Erstelle HTML-Datei
        filename = "output.html"  # Name der Ausgabedatei
        htmlContent = f"""
            <html>
            <head>
                <title>{self.titleInputText}</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                    }}
                    h1 {{
                        color: #333333;
                    }}
                    .author {{
                        font-style: italic;
                    }}
                    .chapter {{
                        font-weight: bold;
                    }}
                    .subchapter {{
                        font-weight: bold;
                    }}
                    .content {{
                        margin-top: 20px;
                    }}
                </style>
            </head>
            <body>
                <h1>{self.titleInputText}</h1>
                <p class="author">Author: {self.authorInputText}</p>
        """

        lines = self.contentInputText.split("\n")
        inParagraph = False

        for line in lines:
            line = line.strip()
            if line.startswith("#"):
                if line.startswith("##"):
                    # Unterüberschrift
                    if inParagraph:
                        htmlContent += "</p>"
                        inParagraph = False
                    htmlContent += f'<h3 class="subchapter">{line[2:].strip()}</h3>'
                else:
                    # Kapitelüberschrift
                    if inParagraph:
                        htmlContent += "</p>"
                        inParagraph = False
                    htmlContent += f'<h2 class="chapter">{line[1:].strip()}</h2>'
            else:
                # Inhalt
                if not inParagraph:
                    htmlContent += "<p>"
                    inParagraph = True
                htmlContent += line.strip() + "<br/>"

        if inParagraph:
            htmlContent += "</p>"

        htmlContent += """
            </body>
            </html>
        """

        with open(filename, "w") as file:
            file.write(htmlContent)

    def createPdf(self):
        # Erstelle PDF-Datei
        self.createHtml()
        #Define path to wkhtmltopdf.exe
        pathToWkhtmltopdf = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
        #Point pdfkit configuration to wkhtmltopdf.exe
        config = pdf.configuration(wkhtmltopdf=pathToWkhtmltopdf)
        inputHtml = "output.html"
        outputPdf = "output.pdf"

        options = {
            'page-size': 'A4',
            'margin-top': '25mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '25mm',
        }

        pdf.from_file(inputHtml, outputPdf, options=options, configuration=config)

    def createWordDocument(self):
        # Erstelle Word-Dokument
        doc = Document()

        # Formatierung anpassen
        titleStyle = doc.styles['Title']
        titleFont = titleStyle.font
        titleFont.size = Pt(18)

        subtitleStyle = doc.styles['Subtitle']
        subtitleFont = subtitleStyle.font
        subtitleFont.size = Pt(12)

        heading1Style = doc.styles['Heading 1']
        heading1Font = heading1Style.font
        heading1Font.size = Pt(16)

        heading2Style = doc.styles['Heading 2']
        heading2Font = heading2Style.font
        heading2Font.size = Pt(14)

        # Inhalte hinzufügen
        titleText = self.titleInputText.replace('\n', '')
        doc.add_heading(titleText, level=1).style = 'Title'
        doc.add_paragraph(f"Author: {self.authorInputText}").style = 'Subtitle'
        doc.add_paragraph(f"Topics: {self.topicsInputText}").style = 'Subtitle'

        lines = self.contentInputText.split("\n")
        inParagraph = False

        for line in lines:
            line = line.strip()
            if line.startswith("#"):
                if line.startswith("##"):
                    # Unterüberschrift
                    if inParagraph:
                        doc.add_paragraph().add_run(line[2:].strip().replace('\n', '')).bold = True
                        inParagraph = False
                    else:
                        doc.add_paragraph().add_run(line[2:].strip().replace('\n', '')).bold = True
                else:
                    # Kapitelüberschrift
                    if inParagraph:
                        doc.add_paragraph().add_run(line[1:].strip().replace('\n', '')).bold = True
                        inParagraph = False
                    else:
                        doc.add_paragraph().add_run(line[1:].strip().replace('\n', '')).bold = True
            else:
                # Inhalt
                if not inParagraph:
                    doc.add_paragraph().add_run(line.strip())
                    inParagraph = True
                else:
                    doc.add_paragraph().add_run(line.strip())

        # Word-Dokument speichern
        outputDocx = "output.docx"
        doc.save(outputDocx)

    def createMarkdown(self):
        # Markdown-Inhalt aus den Inputs erstellen
        author = self.authorInputText
        title = self.titleInputText
        topics = self.topicsInputText
        content = self.contentInputText

        markdownContent = f"# {title}\n\nAuthor: {author}\n\nTopics: {topics}\n\n{content}"

        # Markdown-Datei speichern
        filename = "output.md"
        with open(filename, "w") as file:
            file.write(markdownContent)

    def finish(self):
        # Finish-Funktionalität hier implementieren
        selectedFormats = []
        for checkbox in self.tabWidget.currentWidget().findChildren(QCheckBox):
            if checkbox.isChecked():
                selectedFormats.append(checkbox.text())

        # Beispielhafte Verwendung der ausgewählten Formate
        if "Word Documents" in selectedFormats:
                self.createWordDocument()
        else:
            self.createHtml()
            if "PDF" in selectedFormats:
                self.createPdf()
            if "HTML" in selectedFormats:
                self.createHtml()
            if "Markdown" in selectedFormats:
                self.createMarkdown()
            
    def setAuthorInput(self, text):
        self.authorInputText = text

    def setTitleInput(self, text):
        self.titleInputText = text

    def setTopicsInput(self, text):
        self.topicsInputText = text

    def setContentInput(self, text):
        self.contentInputText = text

    def quitApplication(self):
        app.quit()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle(QStyleFactory.create("Fusion"))  # Set Fusion style for modern look

    font = QFont("Arial", 12)  # Set default font
    app.setFont(font)

    window = MainWindow()
    window.show()
    sys.exit(app.exec_())